import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path("patent_cn")

CONFIG = {
    "font_name_cn": "宋体",
    "font_name_en": "Times New Roman",
    "font_size_pt": 12,
    "line_spacing": 1.5,
    "margin_top_cm": 2.5,
    "margin_bottom_cm": 2.5,
    "margin_left_cm": 2.5,
    "margin_right_cm": 2.5,
}


def configure_document(doc: Document):
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(CONFIG["margin_top_cm"])
    section.bottom_margin = Cm(CONFIG["margin_bottom_cm"])
    section.left_margin = Cm(CONFIG["margin_left_cm"])
    section.right_margin = Cm(CONFIG["margin_right_cm"])

    style = doc.styles["Normal"]
    style.font.name = CONFIG["font_name_en"]
    style.font.size = Pt(CONFIG["font_size_pt"])
    style.element.rPr.rFonts.set(qn("w:eastAsia"), CONFIG["font_name_cn"])
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    style.paragraph_format.line_spacing = CONFIG["line_spacing"]


def add_paragraph_with_style(doc: Document, text: str):
    p = doc.add_paragraph(text)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = CONFIG["line_spacing"]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    return p


def markdown_to_docx(md_path: Path, docx_path: Path):
    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document()
    configure_document(doc)

    for raw in lines:
        line = raw.rstrip()
        if not line.strip():
            doc.add_paragraph("")
            continue

        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
            continue

        if line.startswith("- "):
            add_paragraph_with_style(doc, f"• {line[2:].strip()}")
            continue

        if re.match(r"^\d+\.\s+", line):
            add_paragraph_with_style(doc, line)
            continue

        add_paragraph_with_style(doc, line)

    doc.save(docx_path)


def main():
    tasks = [
        (ROOT / "03_claims/claims_final.md", ROOT / "03_claims/claims_final.docx"),
        (ROOT / "04_spec/spec_draft.md", ROOT / "04_spec/spec_final.docx"),
        (ROOT / "04_spec/abstract.md", ROOT / "04_spec/abstract_final.docx"),
    ]
    for src, dst in tasks:
        markdown_to_docx(src, dst)
        print(f"Generated: {dst}")


if __name__ == "__main__":
    main()

